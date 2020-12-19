import os.path
from excel2json import convert_from_file
import json
import docx
from time import sleep
import sys
from os import system, name 
from datetime import datetime

word_template_name = ""
excel_doc_name = ""
enableVerboseOutput = False # Enable/disable verbose console output

# Verbose console output formatter
def VerbosePrint(message):
    now = datetime.now()
    dt_string = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
    print ("[" + dt_string + "] " + message)	

# Retrieve document names
def GetWordDocName():
    name = input("Name of the Word doc template:")
    return name.strip()
def GetExcelDocName():
    name = input("Name of the Excel doc containing data to parse:")
    return name.strip()

# Set custom bool values
def AskVerbose():
    response = input ("Would you like to enable verbose console output (y/n): ").strip().lower()

    if response == 'y':
        return True
    elif response == 'n':
        return False
    else:
        print("Illegal entry. We will go with No to keep things clean")

# Clear screen
def ClearScreen(): 
    # for windows 
    if name == 'nt': 
        _ = system('cls') 
    # for mac and linux
    else: 
        _ = system('clear') 

# Source file checks
def CheckForExcelDoc():
    if os.path.isfile('contacts.xlsx'):
        return True
    else:
        return False
def CheckForWordTempate():
    if os.path.isfile(word_template_name):
        return True
    else:
        return False

def RemoveAllSpaces(string):
    return string.replace(" ", "")

# Parse Excel sheet to JSON
def ParseToJson():
    if enableVerboseOutput == True:
        VerbosePrint("Attempting to convert Word doc to JSON...")

    try:
        convert_from_file(excel_doc_name)
    except Exception as err:
        print("Error parsing content: " + err + "Exiting program...")
        sleep(3)
        exit(0)

    if enableVerboseOutput == True:
        VerbosePrint("Opening Sheet1.json (default Excel workbook name)")

    input_file = open('Sheet1.json') # Use first (and only) json file to load data from
    contactDict = json.load(input_file)

    GenerateDocuments(contactDict)

# Modify Word doc template and save out changes to new file in generated_docs folder
def GenerateDocuments(contactDict):
    if enableVerboseOutput == True:
        VerbosePrint("Beginning Word doc generation...")

    # { "find this text": "replace with this text" }
    ReplacementDictionary = {"Family Name Here": "FamilyName", "Address Line 1": "AddressLine1", "Address Line 2": "AddressLine2"}

    familyIndex = 0
    familyFileName = ""

    for contact in contactDict:
        doc = docx.Document(word_template_name)

        run = doc.add_paragraph().add_run() # Allow a run to be used on modified text
        style = doc.styles['Normal'] # Normal styling applied
        font = style.font # Accessor to font face
        font.bold = True
        font.name = 'Cavolini' # Set font face
        font.size = docx.shared.Pt(11) # Set font size

        if enableVerboseOutput == True:
            VerbosePrint("Formatting " + contactDict[familyIndex]["FamilyName"])
        
        for i in ReplacementDictionary:
            for p in doc.paragraphs:
                if p.text.find(i) >= 0:
                    p.text = p.text.replace(i, contact[ReplacementDictionary[i]])
                    p.add_run()
        familyIndex += 1

        familyFileName = RemoveAllSpaces(str(contact["FamilyName"]))

        if enableVerboseOutput == True:
            VerbosePrint("Attempting to save " + familyFileName + ".docx")

        try:
            doc.save("generated_docs/" + familyFileName + ".docx")
        except IOError as ex:
            print("Error saving generated document for " + familyFileName +": " + ex)

    print("Process completed. Check generated_docs folder for final outputs.")


if __name__ == "__main__":
    # Ask user for Word doc name
    word_template_name = GetWordDocName()

    # Ask user for Excel file name
    excel_doc_name = GetExcelDocName()

    if CheckForExcelDoc() == False:
        error_response = input("Could not find excel doc with matching name. Would you like to try again? y/n").strip().lower()
        if error_response == 'y':
            ClearScreen()
            os.execv(sys.argv[0], sys.argv)
        elif error_response == 'n':
            exit(0)
        else:
            exit(0)

    if CheckForWordTempate() == False:
        error_response = input("Could not find Word doc with matching name. Would you like to try again? y/n").strip().lower()
        if error_response == 'y':
            ClearScreen()
            os.execv(sys.argv[0], sys.argv)
        elif error_response == 'n':
            exit(0)
        else:
            exit(0)

    enableVerboseOutput = AskVerbose() 

    print("All required files found. Parsing Excel sheet contents and generating Word docs from template")
    sleep(3)
    ParseToJson()

